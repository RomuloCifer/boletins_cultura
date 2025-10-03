import os, re, sys
import pandas as pd
from docx import Document
# from docx2pdf import convert  # opcional (Word no Windows)

# ================== SUPORTE A CAMINHOS (.py e .exe) ==================
def app_dir() -> str:
    if getattr(sys, "_MEIPASS", None):
        return sys._MEIPASS
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def resource_path(*paths):
    base = app_dir()
    return os.path.join(base, *paths)

BASE = os.path.dirname(sys.executable if getattr(sys, "frozen", False) else __file__)

# RESULTADOS e SAÍDA ao lado do app (graváveis)
ARQUIVO_RESULTADOS = os.path.join(BASE, "resultados_boletim.xlsx")
PASTA_SAIDA        = os.path.join(BASE, "boletins_pdf")
os.makedirs(PASTA_SAIDA, exist_ok=True)

# MODELOS (somente leitura) embutidos no EXE ou presentes na pasta do projeto
PASTA_MODELOS = os.path.join(BASE, "Modelos")

# Mapeamento de nível fixo -> modelo correspondente
MAPA_MODELOS = {
    "Lion Stars": "Modelo Boletim - Lion stars.docx",
    "Junior": "Modelo Boletim - Junior.docx",
    "Adultos": "Modelo Boletim - Adolescentes e Adultos.docx",
    "Antigo": "Modelo Boletim - Antigo.docx",
}

# ====== LISTAS DE SUBNÍVEIS ======
ANTIGO_SUBNIVEIS = [
    "High Resolution 4", "High Resolution 5", "High Resolution 6",
    "Basic 5", "Basic 6",
    "New Plus Adult 3",
    "New Plus Adult 2",
    "New Plus Adult 1",
]

ADULTOS_SUBNIVEIS = [
    "Express Pack 1", "Express Pack 2", "Express Pack 3",
    "Inter Teens 1", "Inter Teens 2", "Inter Teens 3",
    "Teen League 1", "Teen League 2", "Teen League 3", "Teen League 4",
    "Mac 1", "Master 2",
]

def _replace_all(texto: str, dados: dict) -> str:
    for k, val in dados.items():
        if pd.isna(val):
            val = ""
        texto = texto.replace(f"<<{k}>>", str(val))
    return texto

def _replace_in_paragraph(par, dados):
    old = "".join(run.text for run in par.runs) or par.text
    new = _replace_all(old, dados)
    if new != old:
        if not par.runs:
            par.add_run()
        par.runs[0].text = new
        for run in par.runs[1:]:
            run.text = ""

def _replace_in_cell(cell, dados):
    if not cell.paragraphs:
        cell.add_paragraph("")
    for p in list(cell.paragraphs):
        _replace_in_paragraph(p, dados)

def substituir_texto(doc: Document, dados: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, dados)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_cell(cell, dados)

def safe_filename(s: str) -> str:
    s = str(s)
    s = re.sub(r'[\\/:*?"<>|]+', "_", s)
    return s.strip()

def gerar_boletins():
    df = pd.read_excel(ARQUIVO_RESULTADOS)

    if "Nome" in df.columns and "Aluno" not in df.columns:
        df = df.rename(columns={"Nome": "Aluno"})

    for _, row in df.iterrows():
        nivel  = str(row.get("Nivel", "")).strip()    # ex.: "High Resolution 5" ou "Inter Teens 2"
        modelo = str(row.get("Modelo", "")).strip()   # "Antigo" quando for um dos antigos

        # Escolha do modelo:
        if modelo == "Antigo" or nivel in ANTIGO_SUBNIVEIS:
            caminho_modelo = os.path.join(PASTA_MODELOS, MAPA_MODELOS["Antigo"])
        elif nivel in ADULTOS_SUBNIVEIS or nivel.startswith("Adultos"):
            caminho_modelo = os.path.join(PASTA_MODELOS, MAPA_MODELOS["Adultos"])
        elif nivel in MAPA_MODELOS:
            caminho_modelo = os.path.join(PASTA_MODELOS, MAPA_MODELOS[nivel])
        else:
            print(f"⚠️ Nível/Modelo sem template: Nivel='{nivel}', Modelo='{modelo}' — pulando {row.get('Aluno', '')}")
            continue

        if not os.path.exists(caminho_modelo):
            print(f"❌ Modelo não encontrado: {caminho_modelo}")
            continue

        dados = {col: row[col] for col in df.columns if col in row.index}
        doc = Document(caminho_modelo)
        substituir_texto(doc, dados)

        nome = safe_filename(dados.get("Aluno", "Aluno"))
        turma = safe_filename(dados.get("Turma", "Turma"))
        nome_docx = f"{nome}_{turma}_{nivel}.docx".replace("/", "-")
        caminho_docx = os.path.join(PASTA_SAIDA, nome_docx)

        doc.save(caminho_docx)

        # Para PDF automático via Word (Windows + Office), descomente:
        # convert(caminho_docx, os.path.join(PASTA_SAIDA, f"{nome}_{turma}_{nivel}.pdf"))

        print(f"✅ Boletim gerado: {caminho_docx}")

if __name__ == "__main__":
    gerar_boletins()
